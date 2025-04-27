# scripts/policy_summarizer.py

import os
import pdfplumber
from docx import Document
import spacy
import numpy as np
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ---------------------- Text Extraction ----------------------
def extract_text(file_path):
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    else:
        raise ValueError("❌ Unsupported file type. Only PDF or DOCX allowed.")

# ---------------------- Summarization ----------------------
@st.cache_resource
def load_spacy_model():
    return spacy.load("en_core_web_sm")

nlp = load_spacy_model()

def spacy_extractive_summary(text, num_sentences=7):
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 20]

    if len(sentences) <= num_sentences:
        return "\n".join(sentences)

    tfidf = TfidfVectorizer(stop_words='english')
    tfidf_matrix = tfidf.fit_transform(sentences)
    similarity_matrix = cosine_similarity(tfidf_matrix)

    sentence_scores = similarity_matrix.sum(axis=1)
    ranked_sentences = [sentences[i] for i in np.argsort(-sentence_scores)[:num_sentences]]

    return "\n".join(ranked_sentences)

# ---------------------- Save Summary ----------------------
def save_summary_to_docx(summary_text, output_filename):
    doc = Document()
    doc.add_heading("Extractive Summary", level=1)
    doc.add_paragraph(summary_text)
    doc.save(output_filename)

# ---------------------- Streamlit UI ----------------------
def run():
    st.title("📄 Policy Document Summarizer")
    st.subheader("Upload your Policy PDF or DOCX")

    uploaded_file = st.file_uploader("Choose a PDF or DOCX file", type=["pdf", "docx"])
    num_sentences = st.slider("Select number of sentences for summary", min_value=3, max_value=20, value=7)

    if uploaded_file:
        # Save uploaded file temporarily
        temp_file_path = os.path.join("temp_upload", uploaded_file.name)
        os.makedirs("temp_upload", exist_ok=True)
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        st.success("✅ File Uploaded Successfully!")
        
        if st.button("Generate Summary"):
            with st.spinner("🔄 Extracting and Summarizing... Please wait..."):
                text = extract_text(temp_file_path)
                summary = spacy_extractive_summary(text, num_sentences=num_sentences)

                st.subheader("📝 Generated Summary:")
                st.text_area("Summary", summary, height=300)

                # Save to output
                output_filename = temp_file_path.replace(".pdf", "_summary.docx").replace(".docx", "_summary.docx")
                save_summary_to_docx(summary, output_filename)

                with open(output_filename, "rb") as f:
                    st.download_button(
                        label="📥 Download Summary as DOCX",
                        data=f,
                        file_name=os.path.basename(output_filename),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            # Clean temporary files
            os.remove(temp_file_path)
            os.remove(output_filename)

if __name__ == "__main__":
    run()
